import io
import pandas as pd
import psycopg2
import streamlit as st
import plotly.express as px
from docx import Document
from docx.shared import Inches

DSN = "dbname=bdl_local user=bdl_user password=mocne_haslo host=localhost port=5432"

def display_variable_name(row):
    # Etykiety pobierane są dynamicznie z bazy (user_label),
    # co eliminuje potrzebę utrzymywania słownika w kodzie.
    return row.get("name")

st.set_page_config(page_title="BaDyL - Lokalny BDL", layout="wide")


@st.cache_resource
def get_conn():
    return psycopg2.connect(DSN)


@st.cache_data(ttl=300)
def get_categories():
    conn = get_conn()
    return pd.read_sql("""
        SELECT DISTINCT category
        FROM variables
        ORDER BY category
    """, conn)["category"].tolist()


@st.cache_data(ttl=300)
def get_all_gminy():
    conn = get_conn()
    return pd.read_sql("""
        SELECT
            g.id,
            g.name,
            g.level,
            g.parentid,
            p.name AS powiat_name
        FROM units g
        LEFT JOIN units p ON p.id = g.parentid AND p.level = 5
        WHERE g.level = 6
        ORDER BY g.name
    """, conn)

@st.cache_data(ttl=300)
def search_units(q):
    conn = get_conn()
    return pd.read_sql("""
        SELECT
            u.id,
            u.name,
            u.level,
            u.parentid,
            p.name AS powiat_name
        FROM units u
        LEFT JOIN units p
            ON p.id = u.parentid
           AND p.level = 5
        WHERE u.name ILIKE %(q)s
           OR u.id ILIKE %(q)s
        ORDER BY u.level, u.name, u.id
        LIMIT 500
    """, conn, params={"q": f"%{q}%"})

@st.cache_data(ttl=300)
def get_units_by_parent(parentid=None):
    conn = get_conn()

    if parentid is None:
        return pd.read_sql("""
            SELECT id, name, level, parentid
            FROM units
            WHERE parentid IS NULL OR level = 0
            ORDER BY level, name
        """, conn)

    return pd.read_sql("""
        SELECT id, name, level, parentid
        FROM units
        WHERE parentid = %(parentid)s
        ORDER BY level, name
    """, conn, params={"parentid": parentid})


@st.cache_data(ttl=300)
def get_wojewodztwa():
    conn = get_conn()
    return pd.read_sql("""
        SELECT id, name, level, parentid
        FROM units
        WHERE level = 2
        ORDER BY name
    """, conn)


@st.cache_data(ttl=300)
def get_powiaty_for_woj(woj_id):
    conn = get_conn()
    prefix = str(woj_id)[:4]

    return pd.read_sql("""
        SELECT id, name, level, parentid
        FROM units
        WHERE level = 5
          AND LEFT(id, 4) = %(prefix)s
        ORDER BY name
    """, conn, params={"prefix": prefix})


@st.cache_data(ttl=300)
def get_gminy_for_powiat(powiat_id):
    conn = get_conn()

    return pd.read_sql("""
        SELECT
            g.id,
            g.name,
            g.level,
            g.parentid,
            p.name AS powiat_name
        FROM units g
        LEFT JOIN units p ON p.id = g.parentid
        WHERE g.level = 6
          AND g.parentid = %(powiat_id)s
        ORDER BY g.name
    """, conn, params={"powiat_id": powiat_id})

def add_unit(unit_id, unit_name, level):
    if "selected_units_dict" not in st.session_state:
        st.session_state.selected_units_dict = {}

    st.session_state.selected_units_dict[str(unit_id)] = {
        "id": str(unit_id),
        "name": unit_name,
        "level": int(level) if level is not None else None
    }


def remove_unit(unit_id):
    if "selected_units_dict" in st.session_state:
        st.session_state.selected_units_dict.pop(str(unit_id), None)

def add_variable(variable_id, variable_name, category, measureunitname):
    if "selected_variables_dict" not in st.session_state:
        st.session_state.selected_variables_dict = {}

    variable_id = int(variable_id)

    st.session_state.selected_variables_dict[variable_id] = {
        "id": variable_id,
        "name": variable_name,
        "category": category,
        "measureunitname": measureunitname
    }


def remove_variable(variable_id):
    if "selected_variables_dict" in st.session_state:
        st.session_state.selected_variables_dict.pop(int(variable_id), None)

def clear_units():
    st.session_state.selected_units_dict = {}


@st.cache_data(ttl=300)
def get_variables(category):
    conn = get_conn()
    return pd.read_sql("""
        SELECT
            id,
            COALESCE(user_label, display_name, subject_name || ' - ' || name, name) AS name,
            category,
            measureunitname
        FROM variables
        WHERE category = %(category)s
        ORDER BY COALESCE(display_name, subject_name || ' - ' || name, name)
    """, conn, params={"category": category})


@st.cache_data(ttl=300)
def search_variables_all(q):
    conn = get_conn()
    return pd.read_sql("""
        SELECT
            id,
            COALESCE(user_label, display_name, subject_name || ' - ' || name, name) AS name,
            category,
            measureunitname
        FROM variables
        WHERE COALESCE(user_label, display_name, subject_name || ' - ' || name, name) ILIKE %(q)s
           OR user_label ILIKE %(q)s
           OR display_name ILIKE %(q)s
           OR subject_name ILIKE %(q)s
           OR name ILIKE %(q)s
           OR CAST(id AS TEXT) ILIKE %(q)s
           OR category ILIKE %(q)s
        ORDER BY category, COALESCE(user_label, display_name, subject_name || ' - ' || name, name)
        LIMIT 300
    """, conn, params={"q": f"%{q}%"})



@st.cache_data(ttl=300)
def get_data(unit_ids, variable_ids, years):
    conn = get_conn()
    return pd.read_sql("""
        SELECT
            d.unit_id,
            u.name AS unit_name,
            d.variable_id,
            COALESCE(v.user_label, v.display_name, v.subject_name || ' - ' || v.name, v.name) AS variable_name,
            v.category,
            d.year,
            d.value,
            v.measureunitname
        FROM data d
        JOIN units u ON u.id = d.unit_id
        JOIN variables v ON v.id = d.variable_id
        WHERE d.unit_id = ANY(%(unit_ids)s)
          AND d.variable_id = ANY(%(variable_ids)s)
          AND d.year = ANY(%(years)s)
        ORDER BY u.name, v.category, v.name, d.year
    """, conn, params={
        "unit_ids": unit_ids,
        "variable_ids": variable_ids,
        "years": years
    })


def to_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="BDL")
    return output.getvalue()


def to_word(df):
    output = io.BytesIO()
    doc = Document()
    doc.add_heading('Eksport danych z BaDyL', 0)

    # Pivotujemy dane, aby lata były w kolumnach
    pivot_df = df.pivot_table(
        index=['unit_name', 'variable_name', 'category'],
        columns='year',
        values='value'
    ).reset_index()
    pivot_df.columns.name = None

    # Pobieramy kolumny lat (liczbowe) i sortujemy je
    years_cols = sorted([c for c in pivot_df.columns if isinstance(c, (int, float))])

    # Obliczamy zmianę procentową między ostatnim a pierwszym rokiem
    if len(years_cols) >= 2:
        first_year = years_cols[0]
        last_year = years_cols[-1]
        pivot_df['Zmiana [%]'] = pivot_df.apply(
            lambda row: ((row[last_year] - row[first_year]) / row[first_year] * 100)
            if row[first_year] != 0 and pd.notna(row[first_year]) and pd.notna(row[last_year])
            else None,
            axis=1
        )
    else:
        pivot_df['Zmiana [%]'] = None

    # Dodajemy tabelę
    t = doc.add_table(rows=1, cols=len(pivot_df.columns))
    t.style = 'Table Grid'

    # Nagłówki
    for i, col in enumerate(pivot_df.columns):
        if i < 3:
            # Pierwsze trzy kolumny (jednostka, wskaźnik, kategoria) mają puste nagłówki wg wzoru
            t.cell(0, i).text = ""
        elif isinstance(col, (int, float)):
            t.cell(0, i).text = str(int(col))
        else:
            t.cell(0, i).text = str(col)

    # Dane
    for _, row in pivot_df.iterrows():
        row_cells = t.add_row().cells
        for i, value in enumerate(row):
            if pd.isna(value):
                row_cells[i].text = ""
            elif isinstance(value, (float, int)) and i >= 3:
                # Formatowanie liczb z przecinkiem jako separatorem dziesiętnym
                row_cells[i].text = f"{value:.2f}".replace('.', ',')
            else:
                row_cells[i].text = str(value)

    doc.save(output)
    return output.getvalue()


if "selected_units_dict" not in st.session_state:
    st.session_state.selected_units_dict = {}

st.title("🎋 BaDyL - Lokalny Mirror BDL")

left, middle, right = st.columns([1.1, 1.2, 2.2])

with left:
    with st.container(border=True):
        st.header("🌍 Jednostki")

        tab_tree, tab_search = st.tabs(["🌳 Drzewo", "🔍 Szukaj"])

    with tab_search:
        unit_query = st.text_input(
            "Szukaj jednostki",
            value="",
            placeholder="np. Cieszyn, Wisła, powiat cieszyński"
        )

        if unit_query:
            units_df = search_units(unit_query)

            if units_df.empty:
                st.info("Brak wyników.")
            else:
                if st.button("Dodaj wszystkie wyniki wyszukiwania", type="secondary"):
                    for _, row in units_df.iterrows():
                        powiat = row.get("powiat_name")
                        if row["level"] == 6 and pd.notna(powiat):
                            add_unit(row["id"], f"{row['name']} | {powiat}", row["level"])
                        else:
                            add_unit(row["id"], row["name"], row["level"])
                    st.rerun()

                st.divider()

                # Ograniczamy listę widoczną do 50, żeby nie spowalniać Streamlit
                for idx, row in units_df.head(50).iterrows():
                    col_a, col_b = st.columns([4, 1])
                    powiat = row.get("powiat_name")

                    if row["level"] == 6 and pd.notna(powiat):
                        label = f"{row['name']} | {powiat} | level {row['level']} | {row['id']}"
                    else:
                        label = f"{row['name']} | level {row['level']} | {row['id']}"

                    with col_a:
                        st.write(label)

                    with col_b:
                        if st.button("Dodaj", key=f"add_search_{row['id']}_{idx}"):
                            if row["level"] == 6 and pd.notna(powiat):
                                add_unit(row["id"], f"{row['name']} | {powiat}", row["level"])
                            else:
                                add_unit(row["id"], row["name"], row["level"])
                            st.rerun()
                if len(units_df) > 50:
                    st.caption(f"...i {len(units_df)-50} innych wyników. Użyj przycisku na górze, aby dodać wszystkie.")

    with tab_tree:
        polska_df = get_units_by_parent(None)
        polska_rows = polska_df[polska_df["level"] == 0]

        col_pol, col_all_gminy = st.columns([1, 1])
        with col_pol:
            if not polska_rows.empty:
                polska = polska_rows.iloc[0]
                if st.checkbox(
                    f"{polska['name']} | {polska['id']}",
                    key=f"tree_polska_{polska['id']}",
                    value=str(polska["id"]) in st.session_state.selected_units_dict
                ):
                    add_unit(polska["id"], polska["name"], polska["level"])
                else:
                    remove_unit(polska["id"])

        with col_all_gminy:
            if st.button("Dodaj WSZYSTKIE gminy w Polsce"):
                all_gminy = get_all_gminy()
                for _, g in all_gminy.iterrows():
                    add_unit(g["id"], f"{g['name']} | {g['powiat_name']}", g["level"])
                st.rerun()

        wojewodztwa = get_wojewodztwa()

        selected_woj_labels = st.multiselect(
            "Województwa",
            [f"{row['name']} | {row['id']}" for _, row in wojewodztwa.iterrows()],
            placeholder="Wybierz jedno lub kilka województw"
        )

        selected_woj_ids = [
            label.split("|")[-1].strip()
            for label in selected_woj_labels
        ]

        if selected_woj_ids:
            if st.button("Dodaj wybrane województwa"):
                for woj_id in selected_woj_ids:
                    woj = wojewodztwa[wojewodztwa["id"] == woj_id].iloc[0]
                    add_unit(woj["id"], woj["name"], woj["level"])
                st.rerun()

            powiaty_all = []
            for woj_id in selected_woj_ids:
                powiaty_all.append(get_powiaty_for_woj(woj_id))

            powiaty = pd.concat(powiaty_all, ignore_index=True) if powiaty_all else pd.DataFrame()

            selected_pow_labels = st.multiselect(
                "Powiaty z wybranych województw",
                [f"{row['name']} | {row['id']}" for _, row in powiaty.iterrows()],
                placeholder="Wybierz jeden lub kilka powiatów"
            )

            selected_pow_ids = [
                label.split("|")[-1].strip()
                for label in selected_pow_labels
            ]

            if selected_pow_ids:
                if st.button("Dodaj wybrane powiaty"):
                    for pow_id in selected_pow_ids:
                        powiat = powiaty[powiaty["id"] == pow_id].iloc[0]
                        add_unit(powiat["id"], powiat["name"], powiat["level"])
                    st.rerun()

                gminy_all = []
                for pow_id in selected_pow_ids:
                    gminy_all.append(get_gminy_for_powiat(pow_id))

                gminy = pd.concat(gminy_all, ignore_index=True) if gminy_all else pd.DataFrame()

                selected_gmina_labels = st.multiselect(
                    "Gminy z wybranych powiatów",
                    [
                        f"{row['name']} | {row['powiat_name']} | {row['id']}"
                        for _, row in gminy.iterrows()
                    ],
                    placeholder="Wybierz jedną lub kilka gmin"
                )

                selected_gmina_ids = [
                    label.split("|")[-1].strip()
                    for label in selected_gmina_labels
                ]

                if selected_gmina_ids:
                    if st.button("Dodaj wybrane gminy"):
                        for gmina_id in selected_gmina_ids:
                            gmina = gminy[gminy["id"] == gmina_id].iloc[0]
                            add_unit(
                                gmina["id"],
                                f"{gmina['name']} | {gmina['powiat_name']}",
                                gmina["level"]
                            )
                        st.rerun()

                if not gminy.empty:
                    if st.button("Dodaj WSZYSTKIE gminy z tych powiatów"):
                        for _, g in gminy.iterrows():
                            add_unit(g["id"], f"{g['name']} | {g['powiat_name']}", g["level"])
                        st.rerun()

    with st.container(border=True):
        st.subheader("📍 Wybrane jednostki")

        num_units = len(st.session_state.selected_units_dict)
        st.metric("Liczba JST", num_units)

        if not st.session_state.selected_units_dict:
            st.info("Wybierz jednostki z powyższego menu.")
        else:
            if st.button("Usuń wszystkie jednostki", key="clear_units_btn", use_container_width=True):
                clear_units()
                st.rerun()

            # Optymalizacja: jeśli jest dużo jednostek, nie pokazujemy przycisków usuwania dla każdej
            if num_units > 20:
                with st.expander(f"Pokaż listę ({num_units})"):
                    for unit_id, unit in st.session_state.selected_units_dict.items():
                        st.write(f"- {unit['name']} (`{unit_id}`)")
            else:
                for unit_id, unit in list(st.session_state.selected_units_dict.items()):
                    col_a, col_b = st.columns([5, 1])
                    with col_a:
                        st.caption(f"{unit['name']}")
                    with col_b:
                        if st.button("❌", key=f"remove_{unit_id}", help=f"Usuń {unit_id}"):
                            remove_unit(unit_id)
                            st.rerun()

    selected_units = list(st.session_state.selected_units_dict.keys())
with middle:
    with st.container(border=True):
        st.header("📊 Wskaźniki")

        if "selected_variables_dict" not in st.session_state:
            st.session_state.selected_variables_dict = {}

        categories = get_categories()
        selected_category = st.selectbox("📂 Kategoria", categories)

    global_indicator_search = st.text_input(
        "Filtruj wskaźniki we wszystkich kategoriach",
        value="",
        placeholder="np. bezrobotni, dochody, żłobki, drogi, 10514"
    )

    indicator_search = st.text_input(
        "Filtruj wskaźniki w aktualnej kategorii",
        value="",
        placeholder="np. uczniowie, przedszkolne, kobiety"
    )

    if global_indicator_search.strip():
        vars_df_visible = search_variables_all(global_indicator_search.strip())
        st.caption(f"Wyniki wyszukiwania ze wszystkich kategorii: {len(vars_df_visible)}")
    else:
        vars_df = get_variables(selected_category)

        if indicator_search.strip():
            q = indicator_search.lower().strip()
            vars_df_visible = vars_df[
                vars_df["name"].str.lower().str.contains(q, na=False)
                | vars_df["id"].astype(str).str.contains(q, na=False)
            ].copy()
        else:
            vars_df_visible = vars_df.copy()

        st.caption(f"Widoczne wskaźniki w kategorii '{selected_category}': {len(vars_df_visible)}")

    var_labels = {
        f"{display_variable_name(row)} | {row['category']} | {row['id']} | {row['measureunitname']}": int(row["id"])
        for _, row in vars_df_visible.iterrows()
    }

    selected_var_labels_to_add = st.multiselect(
        "Wskaźniki z listy",
        list(var_labels.keys()),
        key=f"vars_to_add_{selected_category}_{global_indicator_search}_{indicator_search}"
    )

    col_add_vars, col_add_all = st.columns([1, 1])

    with col_add_vars:
        if st.button("Dodaj wybrane", type="secondary"):
            selected_ids = [var_labels[x] for x in selected_var_labels_to_add]
            for _, row in vars_df_visible[vars_df_visible["id"].isin(selected_ids)].iterrows():
                add_variable(row["id"], row["name"], row["category"], row["measureunitname"])
            st.rerun()

    with col_add_all:
        if st.button("Dodaj wszystkie widoczne", type="secondary"):
            for _, row in vars_df_visible.iterrows():
                add_variable(row["id"], row["name"], row["category"], row["measureunitname"])
            st.rerun()

    with st.container(border=True):
        st.subheader("📈 Wybrane wskaźniki")
        num_vars = len(st.session_state.selected_variables_dict)
        st.metric("Liczba wskaźników", num_vars)

        if not st.session_state.selected_variables_dict:
            st.info("Wyszukaj i dodaj wskaźniki z listy powyżej.")
        else:
            if st.button("Usuń wszystkie wskaźniki", key="clear_vars_btn", use_container_width=True):
                st.session_state.selected_variables_dict = {}
                st.rerun()

            if num_vars > 20:
                with st.expander(f"Pokaż listę ({num_vars})"):
                    for var_id, var in st.session_state.selected_variables_dict.items():
                        st.write(f"- {var['name']} (`{var_id}`)")
            else:
                for var_id, var in list(st.session_state.selected_variables_dict.items()):
                    col_a, col_b = st.columns([5, 1])
                    with col_a:
                        st.caption(f"{var['name']}")
                    with col_b:
                        if st.button("❌", key=f"remove_var_{var_id}", help=f"Usuń {var_id}"):
                            remove_variable(var_id)
                            st.rerun()

    selected_vars = list(st.session_state.selected_variables_dict.keys())

    with st.container(border=True):
        st.header("📅 Okres")
        year_from, year_to = st.slider(
            "Zakres lat",
            min_value=2014,
            max_value=2025,
            value=(2014, 2025)
        )
        years = list(range(year_from, year_to + 1))

    load = st.button("🚀 POBIERZ DANE", type="primary", use_container_width=True)

with right:
    st.header("📋 Wyniki")

    if load:
        if not selected_units:
            st.warning("Wybierz co najmniej jedną jednostkę.")
        elif not selected_vars:
            st.warning("Wybierz co najmniej jeden wskaźnik.")
        else:
            df = get_data(selected_units, selected_vars, years)

            if df.empty:
                st.info("Brak danych dla wybranego zestawu.")
            else:
                tab_table, tab_charts = st.tabs(["📄 Tabela", "📈 Wykresy"])

                with tab_table:
                    st.dataframe(df, use_container_width=True, height=500)

                    col1, col2, col3 = st.columns(3)
                    with col1:
                        csv = df.to_csv(index=False, sep=";", encoding="utf-8-sig")
                        st.download_button(
                            "📥 Pobierz CSV",
                            data=csv,
                            file_name="bdl_export.csv",
                            mime="text/csv",
                            use_container_width=True
                        )
                    with col2:
                        st.download_button(
                            "📥 Pobierz Excel",
                            data=to_excel(df),
                            file_name="bdl_export.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                    with col3:
                        st.download_button(
                            "📥 Pobierz Word",
                            data=to_word(df),
                            file_name="bdl_export.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

                with tab_charts:
                    chart_vars = df["variable_name"].unique()
                    for var in chart_vars:
                        subset = df[df["variable_name"] == var]
                        fig = px.line(
                            subset,
                            x="year",
                            y="value",
                            color="unit_name",
                            title=f"Trend: {var}",
                            markers=True
                        )
                        st.plotly_chart(fig, use_container_width=True)
